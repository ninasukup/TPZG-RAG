import re
import os
import json
import argparse
from pathlib import Path
from typing import List, Dict, Any, Tuple, Optional
from collections import defaultdict
from datetime import datetime

try:
    import camelot
except ImportError:
    camelot = None

from docx import Document
from docx.table import _Table
import fitz 
import pdfplumber
from langchain.text_splitter import RecursiveCharacterTextSplitter
from thefuzz import fuzz

class ChunkingPipeline:
    """
        Features:
    - Processes DOCX and PDF files.
    - Extracts text, tables (including fallback methods), and metadata.
    - Chunks text semantically by sections and then by size.
    - Enriches chunks with metadata: source, author, dates, page numbers, and section hierarchy.
    - Merges small chunks intelligently, combining metadata where appropriate.
    - Provides robust page mapping for PDF chunks.
    """
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200, min_chunk_size: int = 50):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", ". ", " ", ""],
            length_function=len
        )
        self.min_chunk_size = min_chunk_size
        self.section_pattern = re.compile(
            r'^(?P<header>(?:\d+(?:\.\d+)*\s+[A-Za-z].*|'
            r'Chapter\s+\d+|Section\s+\d+|Part\s+[A-Z0-9]+|'
            r'Appendix\s+[A-Z0-9]+))\s*$',
            re.MULTILINE | re.IGNORECASE
        )
        self.doc_id_pattern = re.compile(r'(?:Ref|Proposal|ID|Reference No\.):\s*([A-Z0-9-]+)', re.IGNORECASE)

    def _get_file_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extracts file system and document-internal metadata."""
        file_stat = os.stat(file_path)
        meta = {
            "source_filename": Path(file_path).name,
            "last_modified_date": datetime.fromtimestamp(file_stat.st_mtime).isoformat(),
            "file_size_bytes": file_stat.st_size,
            "author": "Unknown",
            "creation_date": "Unknown",
        }
        ext = Path(file_path).suffix.lower()
        try:
            if ext == ".docx":
                doc = Document(file_path)
                props = doc.core_properties
                meta["author"] = props.author or "Unknown"
                meta["creation_date"] = props.created.isoformat() if props.created else "Unknown"
            elif ext == ".pdf":
                with fitz.open(file_path) as doc:
                    doc_meta = doc.metadata
                    meta["author"] = doc_meta.get("author", "Unknown")
                    meta["creation_date"] = doc_meta.get("creationDate", "Unknown")
        except Exception:
            # File might be corrupt or lack metadata
            pass
        return meta
    
    def _extract_pdf_tables(self, file_path: str, page_num: int) -> List[str]:
        """Extracts tables from a specific PDF page using multiple methods."""
        # Method 1: pdfplumber (primary)
        try:
            with pdfplumber.open(file_path) as pdf:
                page = pdf.pages[page_num - 1]
                tables = page.extract_tables()
                if tables:
                    return ["\n".join([" | ".join(map(str, row)) for row in table if row]) for table in tables]
        except Exception:
            pass # Fallback to next method

        # Method 2: Camelot (fallback, if installed)
        if camelot:
            try:
                # Camelot is sensitive, requires page string
                tables = camelot.read_pdf(file_path, pages=str(page_num), flavor='lattice')
                if tables.n > 0:
                    return [table.df.to_markdown(index=False) for table in tables]
            except Exception:
                pass # Extraction failed
        return []
    
    def load_pdf_per_page(self, file_path: str) -> Dict[int, Dict[str, Any]]:
        """Loads PDF content page by page, including tables."""
        pages_content = defaultdict(lambda: {'text': '', 'tables': []})
        with fitz.open(file_path) as doc:
            for page_num, page in enumerate(doc, 1):
                pages_content[page_num]['text'] = page.get_text().strip()
                # Use our robust table extraction method
                pages_content[page_num]['tables'] = self._extract_pdf_tables(file_path, page_num)
        return dict(pages_content)
    
    def _merge_small_chunks(self, chunks: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Merges small chunks and intelligently combines their metadata."""
        if len(chunks) < 2:
            return chunks
        
        merged_chunks = []
        buffer = chunks[0]

        for i in range(1, len(chunks)):
            next_chunk = chunks[i]
            if len(buffer['content']) < self.min_chunk_size:
                # Merge content
                buffer['content'] += "\n\n" + next_chunk['content']
                # Merge metadata
                # Use a set to avoid duplicate page numbers
                buffer_pages = set(buffer['metadata'].get('page_numbers', []))
                next_pages = set(next_chunk['metadata'].get('page_numbers', []))
                buffer['metadata']['page_numbers'] = sorted(list(buffer_pages.union(next_pages)))
                # Keep the hierarchy of the first (primary) chunk
            else:
                merged_chunks.append(buffer)
                buffer = next_chunk
        
        merged_chunks.append(buffer)
        return merged_chunks
    
    def process_single_document(self, file_path: str, document_id: Optional[str] = None) -> List[Dict[str, Any]]:
        """Processes one file, returning a list of structured chunks."""
        base_metadata = self._get_file_metadata(file_path)
        ext = base_metadata["source_filename"].split('.')[-1]
        all_chunks = []

        if ext == "docx":
            # For DOCX, we process the whole document at once
            doc = Document(file_path)
            full_text = "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])
            base_metadata["document_id"] = document_id or self._extract_document_id(full_text[:2000])
            
            # Since we can't get page numbers, we chunk the whole text
            text_chunks = self._chunk_text_by_section(full_text, base_metadata)
            all_chunks.extend(text_chunks)
            
            # Process tables
            for table in doc.tables:
                table_md = "\n".join([" | ".join([cell.text.strip() for cell in row.cells]) for row in table.rows])
                if table_md.strip():
                    table_meta = base_metadata.copy()
                    table_meta.update({"chunk_type": "table"})
                    all_chunks.append({"content": table_md, "metadata": table_meta})
        
        elif ext == "pdf":
            # --- 4. Improved PDF Chunk-to-Page Mapping ---
            # Process page by page for accurate mapping
            pages_content = self.load_pdf_per_page(file_path)
            if not pages_content: return []
            
            first_page_text = pages_content.get(1, {}).get('text', '')
            base_metadata["document_id"] = document_id or self._extract_document_id(first_page_text)

            for page_num, content in pages_content.items():
                page_text = content.get('text', '')
                page_meta = base_metadata.copy()
                page_meta['page_numbers'] = [page_num]

                if page_text:
                    # Chunk this single page's text
                    page_chunks = self._chunk_text_by_section(page_text, page_meta)
                    all_chunks.extend(page_chunks)

                for table_md in content.get('tables', []):
                    if table_md.strip():
                        table_meta = base_metadata.copy()
                        table_meta.update({"chunk_type": "table", "page_numbers": [page_num]})
                        all_chunks.append({"content": table_md, "metadata": table_meta})
        else:
            print(f"Warning: Unsupported file type '{ext}'. Skipping.")
            return []

        return self._merge_small_chunks(all_chunks)
    
    def _chunk_text_by_section(self, text: str, base_metadata: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Splits text by semantic sections, then chunks recursively."""
        chunks = []
        sections = self.section_pattern.split(text)
        current_hierarchy = base_metadata.get('section_hierarchy', [])
        
        # Process content before the first header
        if sections[0].strip():
            sub_chunks = self.text_splitter.split_text(sections[0])
            for sc in sub_chunks:
                meta = base_metadata.copy()
                meta.update({"chunk_type": "text", "section_hierarchy": current_hierarchy})
                chunks.append({"content": sc, "metadata": meta})
        
        # Process content after each header
        for i in range(1, len(sections), 2):
            header, content = sections[i], sections[i+1]
            if header:
                level = header.strip().split(' ')[0].count('.')
                current_hierarchy = current_hierarchy[:level]
                current_hierarchy.append(header.strip())
            
            if content.strip():
                sub_chunks = self.text_splitter.split_text(content)
                for sc in sub_chunks:
                    meta = base_metadata.copy()
                    meta.update({"chunk_type": "text", "section_hierarchy": current_hierarchy[:]})
                    chunks.append({"content": sc, "metadata": meta})
        return chunks
    
    def _extract_document_id(self, text: str) -> Optional[str]:
        match = self.doc_id_pattern.search(text)
        return match.group(1) if match else None  

def group_files_by_prefix(directory: str) -> Dict[str, List[str]]:
    """Groups files by their common prefix (e.g., 'PROJ-123_tech.pdf')."""
    proposals = defaultdict(list)
    for filename in os.listdir(directory):
        if '_' in filename:
            prefix = filename.split('_')[0]
            proposals[prefix].append(os.path.join(directory, filename))
        else:
            # Handle files with no underscore as their own proposal
            prefix = Path(filename).stem
            proposals[prefix].append(os.path.join(directory, filename))
    return proposals        
        
def process_proposal(proposal_files: List[str], proposal_id: str, pipeline: ChunkingPipeline) -> List[Dict[str, Any]]:
    """Orchestrator to process a list of files for a single proposal."""
    all_proposal_chunks = []
    doc_id = proposal_id
    
    for file_path in proposal_files:
        if not os.path.exists(file_path):
            print(f"Error: File not found: '{file_path}'. Skipping.")
            continue
        
        print(f"  📄 Processing file: {Path(file_path).name}...")
        try:
            file_chunks = pipeline.process_single_document(file_path, document_id=doc_id)
            if not doc_id and file_chunks:
                doc_id = file_chunks[0]['metadata'].get('document_id')
            all_proposal_chunks.extend(file_chunks)
        except Exception as e:
            print(f"  ❌ An error occurred processing {Path(file_path).name}: {e}")
            
    # Final pass to ensure consistent document_id
    for chunk in all_proposal_chunks:
        chunk['metadata']['document_id'] = doc_id
        
    return all_proposal_chunks

def main():
    parser = argparse.ArgumentParser(description="Process and chunk proposal documents for RAG.")
    parser.add_argument("input_dir", type=str, help="Directory containing proposal files (DOCX, PDF).")
    parser.add_argument("--output_file", type=str, default="chunks_output.json", help="Path to save the JSON output file.")
    parser.add_argument("--chunk_size", type=int, default=1000, help="Target size for text chunks.")
    parser.add_argument("--chunk_overlap", type=int, default=200, help="Overlap between consecutive text chunks.")
    args = parser.parse_args()

    if not os.path.isdir(args.input_dir):
        print(f"Error: Input directory not found at '{args.input_dir}'")
        return

    # Initialize the pipeline with CLI arguments
    pipeline = ChunkingPipeline(chunk_size=args.chunk_size, chunk_overlap=args.chunk_overlap)
    
    # Group files into proposals
    proposals_to_process = group_files_by_prefix(args.input_dir)
    all_processed_chunks = {}

    print(f"🔍 Found {len(proposals_to_process)} proposal(s) in '{args.input_dir}'.")

    for proposal_id, file_list in proposals_to_process.items():
        print(f"\n▶️  Processing Proposal: {proposal_id}")
        chunks = process_proposal(file_list, proposal_id, pipeline)
        all_processed_chunks[proposal_id] = chunks
        print(f"✅ Finished proposal {proposal_id}, generated {len(chunks)} chunks.")

    # Save the output
    try:
        with open(args.output_file, 'w', encoding='utf-8') as f:
            json.dump(all_processed_chunks, f, indent=2, ensure_ascii=False)
        print(f"\n💾 Successfully saved all chunks to '{args.output_file}'")
    except Exception as e:
        print(f"\n❌ Error saving output to file: {e}")

if __name__ == "__main__":
    main()